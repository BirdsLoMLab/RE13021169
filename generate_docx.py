#!/usr/bin/env python3
"""Generate a plain-English DOCX version of the LOM Master Formula Reference."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.3)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()  # spacing
    return table

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_heading('Legend of Mushroom\nCombat Formula Reference', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Reverse-Engineered from Game Source Code')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plain English Edition — March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body('')
add_body('This document explains every combat formula in Legend of Mushroom in plain English, with the actual game code included for reference. All formulas are derived directly from the game\'s JavaScript source (457,538 lines of code). Where the code disagrees with community documentation, the code wins.')

doc.add_page_break()

# ============================================================
# SECTION 1: HOW THE GAME DOES MATH
# ============================================================
add_heading('1. How the Game Does Math', 1)

add_body('Before diving into damage formulas, you need to understand how the game handles numbers. LOM uses a custom rounding system called "FixMath" that affects every single calculation in combat. Getting this wrong in a simulator will produce slightly different numbers.')

add_heading('Rounding to 4 Decimal Places', 2)
add_body('Every time the game multiplies or divides two numbers, it rounds the result to exactly 4 decimal places. For example, if a calculation produces 1.23456, the game stores it as 1.2346. This is done using the "round" function:')
add_code('round(x) = floor(10000 * x + 0.5) / 10000')
add_body('Think of it as: multiply by 10,000, round to the nearest whole number, then divide by 10,000.')

add_heading('Rounding Down to Whole Numbers', 2)
add_body('When the game needs a whole number (like final damage), it uses "roundInt" — which first rounds to 4 decimals, then drops everything after the decimal point (always rounds down):')
add_code('roundInt(x) = floor(round(x))')
add_body('For example: roundInt(99.9999) = floor(round(99.9999)) = floor(100.0000) = 100. But roundInt(99.9994) = floor(99.9994) = 99.')

add_note('Why this matters: The game applies roundInt at EVERY multiplication step — not just at the end. A simulator that calculates everything in one step and rounds once at the end will get slightly different results.')

add_heading('The 10,000x Storage Convention', 2)
add_body('Many values in the game are stored as whole numbers that are 10,000 times larger than their actual value. For example, a PvP reduction factor of 25.0 is stored as 250000 and divided by 10,000 at runtime. A shield decay of 40% is stored as 4000. When you see "/ 1e4" or "/ 10000" in the code, this is what\'s happening.')

doc.add_page_break()

# ============================================================
# SECTION 2: THE STAT SYSTEM
# ============================================================
add_heading('2. The Stat (Attribute) System', 1)

add_body('Every unit in combat has a set of attributes identified by number. Here are the most important ones for combat:')

add_heading('Core Stats', 2)
add_table(
    ['ID', 'Internal Name', 'What It Does'],
    [
        ['1001', 'att', 'Your Attack power — the main offensive stat'],
        ['1002', 'hp', 'Your maximum Hit Points'],
        ['1024', 'def', 'Your Defense — subtracted from incoming ATK'],
        ['1060', 'def_coe', 'Defense Coefficient — a hidden multiplier that makes DEF more effective. If you have 0.1 def_coe, your effective DEF is 110% of your base DEF'],
        ['1003', 'att_speed', 'Attack Speed — determines how frequently you attack'],
    ]
)

add_heading('Damage Multipliers', 2)
add_table(
    ['ID', 'Internal Name', 'What It Does'],
    [
        ['1039', 'att_dam', 'Basic ATK multiplier — scales your normal attack damage'],
        ['1032', 'double_hit_dam', 'Combo multiplier — scales your combo (double hit) damage'],
        ['1033', 'counter_dam', 'Counter multiplier — scales your counter-attack damage'],
        ['1045', 'skill_dam_extra', 'Skill damage multiplier — scales your skill damage'],
        ['1040', 'partner_dam', 'Pal damage multiplier — scales your pal\'s damage'],
        ['1047', 'partner_dam_extra', 'Pal damage extra — additional pal damage scaling from the player'],
    ]
)

add_heading('Damage Resistances', 2)
add_table(
    ['ID', 'Internal Name', 'What It Does'],
    [
        ['1018', 'att_resist', 'Basic ATK resistance — reduces incoming normal attacks'],
        ['1034', 'double_hit_def', 'Combo resistance — reduces incoming combo damage'],
        ['1035', 'counter_def', 'Counter resistance — reduces incoming counter damage'],
        ['1019', 'skill_resist', 'Skill resistance — reduces incoming skill damage'],
        ['1020', 'partner_resist', 'Pal resistance — reduces incoming pal damage'],
        ['1021', 'resist', 'DMG Resistance — a general reduction applied to most damage'],
    ]
)

add_heading('Critical Hit Stats', 2)
add_table(
    ['ID', 'Internal Name', 'What It Does'],
    [
        ['1004', 'crit_rate', 'Your chance to land a critical hit'],
        ['1005', 'crit_dam', 'Your critical damage multiplier'],
        ['1006', 'crit_def', 'Reduces incoming crit damage (minimum 0.5 = 50%)'],
        ['1037', 'skill_crit_rate', 'Chance for a skill to critically hit (separate from normal crit)'],
        ['1038', 'skill_crit_dam', 'Bonus damage when a skill crits'],
        ['1065', 'ignore_crit_rate', 'Reduces the attacker\'s effective crit chance'],
    ]
)

add_heading('Other Combat Stats', 2)
add_table(
    ['ID', 'Internal Name', 'What It Does'],
    [
        ['1016', 'double_hit', 'Chance to trigger a combo (double hit) after a normal attack'],
        ['1017', 'counter', 'Chance to counter-attack when hit'],
        ['1007', 'hit', 'Accuracy — reduces your chance of missing'],
        ['1008', 'miss', 'Evasion — increases your chance of dodging'],
        ['1081', 'total_dam_add', 'Total DMG Bonus — adds a final damage multiplier'],
        ['1082', 'total_dam_def', 'Total DMG Resistance — opposes Total DMG Bonus'],
        ['1051', 'shield_hp_extra', 'Bonus shield HP (percentage increase)'],
        ['1046', 'boss_dam', 'Extra damage dealt to boss-type enemies'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 3: DAMAGE FORMULAS
# ============================================================
add_heading('3. How Damage Is Calculated', 1)

add_heading('Step 1: The Base Damage Number', 2)
add_body('All damage in LOM starts the same way. The game takes your ATK and subtracts the enemy\'s effective DEF:')
add_code('base_raw = max(roundInt(ATK - DEF * (1 + DEF_COE)), 1)')
add_body('In plain English: take your Attack, subtract the enemy\'s Defense (boosted by their Defense Coefficient), and round down. The result is always at least 1 — you can never deal zero base damage.')

add_note('IMPORTANT: Community docs say the formula is just "ATK - DEF". The code actually uses "ATK - DEF * (1 + DEF_COE)". The Defense Coefficient (def_coe) is a stat that makes Defense more effective. If someone has a def_coe of 0.2, their effective Defense is 120% of their listed value.')

add_heading('3.1 Basic ATK Damage', 2)
add_body('A normal attack multiplies the base damage by your Basic ATK Multiplier (att_dam), then reduces it by the enemy\'s Basic ATK Resistance (att_resist):')
add_code('damage = roundInt(base_raw * round(ATT_DAM * round(1 - ATT_RESIST)))')
add_body('Notice that the multiplier and resistance are combined into one calculation before being applied to the base damage. This means the game rounds the product of "multiplier times (1 minus resistance)" first, then multiplies by base damage.')
add_body('After this, DMG Resistance is applied (see Step 2 below). Then, if the attack was a critical hit, the crit multiplier is applied (see Section 4).')
add_body('The final result is always at least 1 damage.')

add_heading('3.2 Combo (Double Hit) Damage', 2)
add_body('When a combo triggers, it works slightly differently from a basic attack:')
add_code('damage = roundInt(roundInt(base_raw * DOUBLE_HIT_DAM) * round(1 - DOUBLE_HIT_DEF))')
add_body('Here, the combo multiplier is applied to the base damage FIRST, rounded, and THEN the combo resistance is applied. This is a subtle but important difference from basic ATK, where multiplier and resistance are combined before applying to the base.')

add_heading('3.3 Counter Damage', 2)
add_body('Counter-attacks work the same way as combos, but with counter-specific stats:')
add_code('damage = roundInt(roundInt(base_raw * COUNTER_DAM) * round(1 - COUNTER_DEF))')
add_body('Same structure: multiply base by counter multiplier, round, then apply counter resistance.')

add_heading('3.4 Skill Damage', 2)
add_body('Skill damage follows the same pattern as combo/counter:')
add_code('damage = roundInt(roundInt(base_raw * SKILL_DAM_EXTRA) * round(1 - SKILL_RESIST))')
add_body('Some skills have more complex behavior (see the Skill Crit section below), but the base calculation follows this structure.')

add_heading('Step 2: DMG Resistance (calHurt)', 2)
add_body('After the type-specific damage is calculated, the game applies a general DMG Resistance layer. This is a separate stat from the type-specific resistances above:')
add_code('damage = roundInt(roundInt(damage * round(1 - RESIST)) * round(1 - PVE_RESIST))')
add_body('"RESIST" here is attribute 1021 (DMG Resistance). In PvE, there\'s also a PvE-specific resistance. The result is always at least 1.')
add_body('There\'s also a PvE damage bonus that\'s applied before this: damage = roundInt(damage * round(1 + PVE_DAM)).')

add_heading('3.5 Pal Damage', 2)
add_body('Pal damage has some unique properties:')
add_body('1. Pals use their PARENT PLAYER\'s ATK stat, not their own. Your pal hits harder when your character is stronger.')
add_body('2. The pal\'s damage multiplier combines two stats: the pal\'s own "partner_dam" and the player\'s "partner_dam_extra".')
add_body('3. Instead of basic ATK resistance, the target\'s "partner_resist" (Pal Resistance) is used.')
add_body('4. The Suppress/Inspire system (see Section 10) can modify this resistance.')
add_code('pal_mult = round(PARTNER_DAM * PARENT_PARTNER_DAM_EXTRA)\n'
         'resistance = after_suppress_inspire(PARTNER_RESIST)\n'
         'damage = roundInt(base_raw * round(pal_mult * round(1 - resistance)))')

doc.add_page_break()

# ============================================================
# SECTION 4: CRITICAL HITS
# ============================================================
add_heading('4. Critical Hit System', 1)

add_body('LOM has TWO separate crit systems: normal crit and skill crit. They work differently.')

add_heading('4.1 Normal Crit', 2)
add_body('Every normal attack, combo, and counter rolls for a hit outcome. There are three possibilities: miss, normal hit, or critical hit.')

add_heading('Miss Chance', 3)
add_body('First, the game calculates your chance to miss. It takes the target\'s Evasion and subtracts your Accuracy:')
add_code('raw_evasion = max(round(MISS - HIT), 0)')
add_body('This raw evasion is then put through a "diminishing returns" curve that makes very high evasion less effective:')
add_code('corrected_evasion = round((100 * raw_evasion)^0.9 / 100)')
add_body('The 0.9 exponent means that doubling your evasion does NOT double your miss chance — there are diminishing returns. In PvP, the miss chance is capped at 80% no matter how much evasion you stack.')

add_heading('Crit Chance', 3)
add_body('Your effective crit rate is simply your Crit Rate minus the target\'s Ignore Crit Rate:')
add_code('effective_crit = max(CRIT_RATE - IGNORE_CRIT_RATE, 0)')
add_body('The final probabilities are:')
add_body('  - Miss chance = final_evasion')
add_body('  - Normal hit chance = (1 - evasion) * (1 - effective_crit)')
add_body('  - Crit chance = (1 - evasion) * effective_crit')

add_heading('Crit Damage Multiplier', 3)
add_body('When you crit, your damage is multiplied by:')
add_code('crit_mult = max(1.5, round(CRIT_DAM / max(0.5, CRIT_DEF)))')
add_body('In plain English: divide your Crit Damage by the target\'s Crit Defense. The result is at least 1.5x (you always do at least 50% more damage on a crit). The target\'s Crit Defense can\'t go below 0.5 (50%).')

add_heading('4.2 Skill Crit (Separate System)', 2)
add_body('Skills have their own crit system that\'s completely separate from normal crit:')
add_body('  - Skill crit rate = SKILL_CRIT_RATE (a separate stat)')
add_body('  - Skill crit does NOT check against Ignore Crit Rate')
add_body('  - The damage formula is different:')
add_code('skill_crit_damage = roundInt(Math.pow(roundInt(damage * round(1 + SKILL_CRIT_DAM)), 0.98))')
add_body('This means: multiply the skill damage by (1 + your Skill Crit Damage stat), then raise the ENTIRE result to the power of 0.98.')

add_note('Community docs say the exponent is on (1 + SKILL_CRIT_DAM) alone — i.e., Skill * (1+SCRIT)^0.98. The code shows it\'s on the whole product: (Skill * (1+SCRIT))^0.98. These give different results. The code is correct.')

add_body('The 0.98 exponent is a slight reduction that prevents skill crits from being too dominant at high damage values. It means that as your raw skill crit damage gets higher, you get slightly less than proportional returns.')

doc.add_page_break()

# ============================================================
# SECTION 5: PVP SYSTEM
# ============================================================
add_heading('5. PvP System', 1)

add_heading('5.1 PvP Damage Reduction', 2)
add_body('In PvP, all damage is dramatically reduced by a factor that depends on the average level of both players:')
add_code('avg_level = roundInt((player1_level + player2_level) / 2)\n'
         'injuryReduce = configLevel[avg_level].pvp_injury_reduce / 10000')
add_body('The injuryReduce value gets very large at high levels (around 25x or more), meaning PvP damage is roughly 1/25th of PvE damage. This is why PvP fights last much longer than PvE battles.')
add_body('The reduction is applied as simple division at the very end:')
add_code('final_damage = max(roundInt(damage / injuryReduce), 1)')
add_body('No matter how large the PvP factor, you always deal at least 1 damage.')

add_heading('5.2 Shield Decay in PvP', 2)
add_body('Shields are weaker in PvP. Their HP is multiplied by a "shield decay" factor:')
add_code('shieldDecay = 0.4 (40%)')
add_body('This means shields in PvP only have 40% of their normal value. A shield that would give 100,000 HP in PvE only gives 40,000 in PvP.')

add_note('Shield and heal decay are GLOBAL constants — they do NOT change with player level. Only the damage reduction factor (injuryReduce) is level-dependent.')

add_heading('5.3 Heal Decay in PvP', 2)
add_body('Healing is also reduced in PvP:')
add_code('treatDecay = 0.3 (30%)')
add_body('All healing in PvP is reduced to 30% of its normal value. This includes both active heals and passive HP recovery.')

doc.add_page_break()

# ============================================================
# SECTION 6: HP-BASED DAMAGE
# ============================================================
add_heading('6. HP-Based Damage', 1)

add_body('Some skills deal damage based on a percentage of the target\'s HP. These work very differently from normal damage and bypass several steps in the normal pipeline.')

add_heading('How It Works', 2)
add_body('1. Calculate: HP value * skill percentage')
add_body('2. Multiply UP by the PvP factor (making it temporarily much larger)')
add_body('3. Clamp the result between a minimum and maximum based on your basic ATK damage')
add_body('4. Send this inflated number as damage')
add_body('5. At the damage application step, divide back DOWN by the PvP factor')

add_code('hp_dmg = roundInt(hp_value * skill_percent)\n'
         'hp_dmg = roundInt(hp_dmg * pvp_factor)          // multiply UP\n'
         'base_atk = roundInt(max(roundInt(ATK - DEF*(1+DEF_COE)), 1) * ATT_DAM)\n'
         'min_dmg = roundInt(base_atk * limit_min)         // e.g. 0.8x\n'
         'max_dmg = roundInt(base_atk * limit_max)         // e.g. 50x, 100x, or 2000x\n'
         'hp_dmg = clamp(hp_dmg, min_dmg, max_dmg)\n'
         'final = max(roundInt(hp_dmg / pvp_factor), 1)    // divide BACK DOWN')

add_body('Why the multiply-then-divide trick? It ensures the clamping (min/max) happens against values that make sense relative to your raw ATK damage, not the PvP-reduced values.')

add_heading('Clamp Limits', 2)
add_body('The min/max clamp values come from the skill configuration. Typical values:')
add_body('  - Minimum: 0.8x your basic ATK damage (you always deal at least 80% of a normal attack)')
add_body('  - Maximum for current HP skills: 50x your basic ATK damage')
add_body('  - Maximum for max HP skills: 100x your basic ATK damage')
add_body('  - Maximum for pal HP skills: 2000x your basic ATK damage')

add_note('HP-based damage skips DMG Resistance (resist, ID 1021) — it bypasses calHurt(). However, Total DMG Bonus/RES DOES apply because HP-based damage passes through healthTarget() with HealthType.Hurt, which is in NeedAddDamHurtList. After Total DMG, it goes through PvP reduction and shield absorption.')

doc.add_page_break()

# ============================================================
# SECTION 7: BLEED DAMAGE
# ============================================================
add_heading('7. Bleed Damage', 1)

add_body('The bleed system is more complex than most players realize. There are 8 different types of bleed, each with its own damage formula:')

add_table(
    ['Type', 'Damage Source', 'Special Properties'],
    [
        ['0', 'ATK-based (basic multiplier)', 'Can trigger Giant Slayer buff'],
        ['1', 'Current HP percentage', 'Scales with PvP factor'],
        ['2', 'ATK-based (skill multiplier)', 'Can skill crit (0.98 exponent)'],
        ['3', 'ATK-based (basic + ATK resist)', 'Can normal crit'],
        ['4', 'ATK-based (combo multiplier)', 'Can normal crit'],
        ['5', 'ATK-based (counter multiplier)', 'Can normal crit'],
        ['6', 'Max HP percentage', 'Scales with PvP factor'],
        ['10', 'Target/caster attribute-based', 'Scales with PvP factor'],
    ]
)

add_body('ATK-based bleeds (types 0, 2, 3, 4, 5) all start with the standard "ATK minus effective DEF" base and apply their respective damage multipliers and resistances. HP-based bleeds (types 1, 6, 10) simply take a percentage of the relevant HP value.')

add_body('Bleed DOTs (damage over time) split their total damage evenly across multiple ticks.')

doc.add_page_break()

# ============================================================
# SECTION 8: SHIELDS
# ============================================================
add_heading('8. Shield System', 1)

add_heading('Shield Creation', 2)
add_body('When a shield is created, its HP is calculated like this:')
add_body('1. Start with a base value (from an attribute, ATK-DEF calculation, or HP difference)')
add_body('2. Multiply by the skill\'s shield percentage')
add_body('3. Multiply by (1 + your Shield HP Extra stat) — this is a percentage bonus')
add_body('4. In PvP, multiply by shield decay (0.4 = 40%)')
add_code('shield_hp = roundInt(base * skillPar)\n'
         'shield_hp = roundInt(shield_hp * round(1 + SHIELD_HP_EXTRA))\n'
         'shield_hp = roundInt(shield_hp * 0.4)   // PvP only')

add_heading('How Shields Absorb Damage', 2)
add_body('When you take damage with an active shield:')
add_body('1. The PvP reduction is applied to the damage FIRST')
add_body('2. Then the shield absorbs as much as it can')
add_body('3. Any remaining damage goes through to your HP')
add_body('4. If you have multiple shields, they\'re checked one at a time')
add_body('5. Damage fully overflows — when a shield breaks, the leftover damage isn\'t wasted')

add_note('Some shields are immune to PvP decay (marked with a special flag). These keep their full value even in PvP.')

doc.add_page_break()

# ============================================================
# SECTION 9-10: PIERCE/BLOCK & INSPIRE/SUPPRESS
# ============================================================
add_heading('9. Pierce and Block', 1)

add_body('Pierce and Block don\'t directly modify damage numbers. Instead, they modify the RESISTANCE value that\'s used in the damage formula.')

add_heading('How It Works', 2)
add_body('On each attack, the game rolls to see if Pierce or Block triggers:')
add_body('  - Pierce can trigger if your Armor Pen > target\'s Ignore Armor Pen')
add_body('  - Block can trigger if target\'s Block > your Ignore Block')
add_body('  - They\'re mutually exclusive — only one can happen per attack')
add_body('  - Pierce is checked first (has priority)')

add_body('If Pierce triggers: the target\'s resistance is REDUCED, meaning they take more damage')
add_body('If Block triggers: the target\'s resistance is INCREASED, meaning they take less damage')
add_code('Pierce: resistance -= min(0.5, (armor_pen - ignore_pen) / 10000)\n'
         'Block:  resistance += min(0.5, (block - ignore_block) / 10000)')
add_body('The change is capped at 0.5 (50%) in either direction.')

add_heading('10. Pal Inspire and Suppress', 1)

add_body('Inspire and Suppress work identically to Pierce and Block, but they only affect Pal Resistance (the resistance against pal damage).')
add_body('  - Suppress reduces the target\'s pal resistance (pal deals more damage)')
add_body('  - Inspire increases the target\'s pal resistance (pal deals less damage)')
add_body('Same cap of 0.5, same mutually exclusive check, same priority (suppress first).')

doc.add_page_break()

# ============================================================
# SECTION 11: STUN & CONTROL
# ============================================================
add_heading('11. Stun and Control', 1)

add_heading('Stun Chance', 2)
add_body('The chance to stun works like evasion — it uses the same diminishing returns curve:')
add_code('effective_stun = max(0, round(VERTIGO - VERTIGO_DEF))\n'
         'corrected = round((100 * effective_stun)^0.9 / 100)')
add_body('Your stun rate minus the target\'s stun defense gives the raw chance, which is then put through the 0.9 power curve. High stun rates have diminishing returns.')

add_heading('Stun Duration', 2)
add_body('When a stun lands, its duration is:')
add_code('duration = VERTIGO_TIMES * round(1 - VERTIGO_RES)')
add_body('VERTIGO_TIMES is the base stun duration multiplier, and VERTIGO_RES reduces it. If you have 30% stun duration reduction, the stun lasts 70% as long.')

add_heading('Knockup (Launch)', 2)
add_body('Knockup uses a simpler formula — no diminishing returns curve:')
add_code('probability = round(SUSPEND - SUSPEND_DEF)')
add_body('It\'s a straight subtraction. This means knockup scales linearly, unlike stun.')

doc.add_page_break()

# ============================================================
# SECTION 12: IGNORE MECHANICS
# ============================================================
add_heading('12. Ignore Mechanics', 1)

add_body('All "ignore" stats in LOM work the same way: they directly subtract from the corresponding rate. This is confirmed for every ignore type in the game.')
add_code('effective_rate = max(rate - ignore_rate, 0)')
add_body('For example: if you have 120% counter rate and the enemy has 40% Ignore Counter, your effective counter rate is 80% (not 72% — it\'s subtraction, not multiplication).')
add_body('The result can\'t go below 0 — you can\'t have a negative rate.')

add_table(
    ['Mechanic', 'Your Rate Stat', 'Enemy\'s Ignore Stat'],
    [
        ['Crit', 'crit_rate', 'ignore_crit_rate'],
        ['Combo', 'double_hit', 'ignore_double_hit'],
        ['Counter', 'counter', 'ignore_counter'],
        ['Armor Pen', 'armor_penetration', 'ignore_armor_penetration'],
        ['Block', 'block', 'ignore_block'],
        ['Pal Inspire', 'partner_inspire', 'ignore_partner_inspire'],
        ['Pal Suppress', 'partner_suppress', 'ignore_partner_suppress'],
        ['HP Steal', 'hpsteal_rate', 'hpsteal_res'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 13: BUFF DAMAGE MODIFIERS
# ============================================================
add_heading('13. Buff Damage Modifiers', 1)

add_body('Several buff types modify damage after the base calculation. These are applied in the attack execution code, after the base damage formula but before PvP reduction.')

add_heading('13.1 Extra Damage (BuffExtraDamage)', 2)
add_body('Multiplies your damage by a bonus factor. Three variants:')
add_body('  - Type 0: Fixed percentage bonus — damage * (1 + bonus%)')
add_body('  - Type 1: HP loss scaling — the more HP you\'ve lost, the more bonus damage you get. Formula: damage * (1 + (missing HP / max HP) * bonus%)')
add_body('  - Type 2: Same as Type 1, but considers certain HP-modification buffs when calculating "current HP"')

add_heading('13.2 Giant Slayer (BuffGiantSlayer)', 2)
add_body('Gives bonus damage when fighting enemies with more HP than you:')
add_body('  - Only activates if the target\'s max HP exceeds your max HP')
add_body('  - Bonus scales with the HP difference percentage')
add_body('  - Capped separately for boss targets vs. normal units')
add_code('hp_diff_percent = ceil((target_HP - your_HP) / your_HP * 100)\n'
         'bonus = hp_diff_percent * extraDam_per_percent\n'
         'bonus = min(bonus, cap)\n'
         'damage = damage * (1 + bonus / 10000)')

add_heading('13.3 Fragile Effect (BuffSkillFragileAdd)', 2)
add_body('Adds a flat amount of bonus damage (additive, not multiplicative):')
add_body('  - Type 0: Bonus = attacker\'s [specific stat] * skill parameter')
add_body('  - Type 1: Bonus = target\'s current HP * skill parameter')
add_body('This bonus is ADDED to the damage, not multiplied. For example, if your base damage is 5000 and Fragile adds 1000, you deal 6000.')

add_heading('13.4 Total DMG Bonus / Total DMG RES', 2)
add_body('This is a universal final damage multiplier applied through SkillRunner.healthTarget() to ALL 13 offensive damage types (normal, crit, combo, counter, bleed, real damage, spirit-to-player, shared damage, and return/reflect damage). It is subtractive:')
add_code('multiplier = max(1 + TOTAL_DAM_ADD - TOTAL_DAM_DEF, 0.20)')
add_body('Your Total DMG Bonus minus the enemy\'s Total DMG Resistance, added to 1. The result has a floor of 0.20 (20%) — damage can never be reduced below 20% by this system.')
add_body('This is applied AFTER all buff modifiers (Fragile, Extra Damage, Giant Slayer, boss damage) and BEFORE damage is queued for application. It is the last multiplicative layer before PvP reduction.')

add_note('Total DMG Bonus/RES is SEPARATE from DMG Resistance (resist, ID 1021). Both apply, but at different pipeline stages: DMG Resistance is applied early during calHurt(), while Total DMG is applied late in healthTarget(). They stack multiplicatively. Additionally, BuffVampire (life steal) uses the same formula independently for heal calculations.')

doc.add_page_break()

# ============================================================
# SECTION 14: FULL DAMAGE PIPELINE
# ============================================================
add_heading('14. The Complete Damage Pipeline', 1)

add_body('Here is exactly what happens when damage is dealt, in order:')

add_heading('Attack Phase (calculating the number)', 2)
add_body('1. Calculate base damage: ATK - DEF * (1 + DEF_COE), minimum 1')
add_body('2. Apply the type-specific multiplier (ATT_DAM, DOUBLE_HIT_DAM, COUNTER_DAM, or SKILL_DAM_EXTRA)')
add_body('3. Apply the type-specific resistance (att_resist, double_hit_def, counter_def, or skill_resist)')
add_body('4. Apply DMG Resistance (resist, ID 1021) and PvE bonuses via calHurt')
add_body('5. If critical hit: multiply by max(1.5, CRIT_DAM / CRIT_DEF)')
add_body('6. Apply Boss Damage bonus if target is a boss')
add_body('7. Add Fragile Effect bonus damage (flat addition)')
add_body('8. Apply Extra Damage multiplier')
add_body('9. Apply Giant Slayer HP-based bonus')
add_body('10. Apply recordDamage bonus: damage * (1 + recordDamage[skillId] / 10000) — cumulative per-skill bonus from USE_SKILL_ADD buffs')
add_body('11. Apply counterDamage multiplier: damage * skill.counterDamage — per-skill config multiplier (default 1.0)')

add_heading('Delivery Phase (healthTarget)', 2)
add_body('12. Total DMG Bonus/RES: damage * max(1 + total_dam_add - total_dam_def, 0.20) — applied to all 13 damage types when attacker != target')

add_heading('Application Phase (applying the number)', 2)
add_body('13. Skip if battle is transitioning between phases')
add_body('14. PvP reduction: divide by injuryReduce (minimum 1 damage)')
add_body('15. Season PvE bonus (if applicable): damage * (1 + seasonPveDamAdd), team 1 only, season PvE chapters')
add_body('16. Shield absorption: shields absorb as much as they can, remainder passes through')
add_body('17. Block absorption: block buffs absorb remaining damage')
add_body('18. Subtract remaining damage from HP')
add_body('19. Death prevention checks: Time Reversal > Remake HP > Immune Death')
add_body('20. If prevented from dying: HP stays at 1')
add_body('21. Accumulate total damage dealt')
add_body('22. Check HP change triggers')
add_body('23. If HP <= 0: unit dies')

add_heading('For Healing', 2)
add_body('1. Multiply heal by treatDecay (0.3 in PvP = 30%)')
add_body('2. Apply any REDUCE_HEAL debuffs')
add_body('3. Add to HP, capped at max HP')

doc.add_page_break()

# ============================================================
# SECTION 15: CONFIG CONSTANTS
# ============================================================
add_heading('15. Game Config Constants', 1)

add_body('These are hardcoded values in the game that control combat behavior:')

add_table(
    ['Constant', 'Raw Value', 'Actual Value', 'What It Controls'],
    [
        ['miss_correct', '9000', '0.9', 'The exponent in the evasion diminishing returns curve'],
        ['vertigo_correct', '9000', '0.9', 'The exponent in the stun diminishing returns curve'],
        ['shield_correct', '4000', '0.4 (40%)', 'How much shield HP is reduced in PvP'],
        ['hp_recovery_correct', '3000', '0.3 (30%)', 'How much healing is reduced in PvP'],
        ['battle_up_limit', '8000', '0.8 (80%)', 'Maximum evasion chance in PvP'],
        ['total_damage_add_down_limit', '2000', '0.2 (20%)', 'Minimum Total DMG multiplier (floor)'],
    ]
)

add_note('Shield and heal decay are GLOBAL constants — the same at every level. Only the PvP damage reduction factor (injuryReduce) changes with level.')

doc.add_page_break()

# ============================================================
# SECTION 16: DISCREPANCIES
# ============================================================
add_heading('16. Where Community Docs Are Wrong', 1)

add_body('These are confirmed differences between what the game code actually does and what community documentation (such as Yuko\'s PDF) says:')

add_heading('1. Defense Coefficient is Missing', 2)
add_body('Community docs: ATK - DEF')
add_body('Actual code: ATK - DEF * (1 + DEF_COE)')
add_body('The def_coe stat exists and affects ALL damage formulas. If it\'s non-zero, effective defense is higher than the listed DEF value.')

add_heading('2. Skill Crit Exponent is Misplaced', 2)
add_body('Community docs: Skill * (1 + SkillCritDMG)^0.98 — exponent on the multiplier')
add_body('Actual code: (Skill * (1 + SkillCritDMG))^0.98 — exponent on the whole product')
add_body('This produces different results, especially at high damage values.')

add_heading('3. Total DMG Bonus/RES IS Universal (Confirmed)', 2)
add_body('Community docs: Applied as a "final layer" to all damage')
add_body('Actual code: CONFIRMED — applied universally via SkillRunner.healthTarget() to all 13 damage types in NeedAddDamHurtList (normal, crit, combo, counter, bleed, real damage, shared damage, spirit-to-player, return/reflect).')
add_body('This is indeed a final multiplicative layer, applied after all buff modifiers and before PvP reduction. Community understanding was correct on this point.')

add_heading('4. Total DMG Floor is 20%', 2)
add_body('Community docs: Floor unknown')
add_body('Actual code: Floor = 0.20 (20%). Damage cannot be reduced below 20% by the Total DMG system.')

add_heading('5. Pal Uses Player\'s ATK', 2)
add_body('Community docs: Unclear which ATK is used')
add_body('Actual code: Pal damage always uses the PARENT PLAYER\'s ATK stat, not the pal\'s own ATK.')

add_heading('6. Shield/Heal Decay are Fixed', 2)
add_body('Community docs: Not specified')
add_body('Actual code: Shield = 40%, Heal = 30%. These never change regardless of player level.')

add_heading('7. Rounding Matters', 2)
add_body('Community docs: Simple formulas without rounding')
add_body('Actual code: 10+ rounding operations per damage calculation. Each intermediate result is rounded to 4 decimals or to an integer before the next step. A simulator that skips these rounding steps will produce slightly different numbers.')

add_heading('8. Pierce/Block Modify Resistance, Not Damage', 2)
add_body('Community docs: Often described as direct damage modification')
add_body('Actual code: Pierce reduces the target\'s resistance value; Block increases it. The modified resistance then affects the damage formula indirectly.')

add_heading('9. Evasion Uses Diminishing Returns', 2)
add_body('Community docs: Not detailed')
add_body('Actual code: Evasion uses a power curve with exponent 0.9, creating diminishing returns at high values. PvP cap is 80%.')

add_heading('10. HP-Based Damage Skips DMG Resistance but NOT Total DMG', 2)
add_body('Community docs: Unclear')
add_body('Actual code: HP-based damage is NOT affected by DMG Resistance (resist, ID 1021) — it bypasses calHurt(). However, Total DMG Bonus/RES DOES apply because HP-based damage calls healthTarget() with HealthType.Hurt, which IS in NeedAddDamHurtList. It also goes through PvP reduction, clamping, and shield absorption.')

# ============================================================
# SECTION 17: RECENTLY DISCOVERED MECHANICS
# ============================================================
doc.add_page_break()
add_heading('17. Recently Discovered Mechanics', 1)

add_heading('17.1 Record Damage (Per-Skill Cumulative Bonus)', 2)
add_body('The recordDamage system is a per-skill damage accumulator driven by USE_SKILL_ADD buffs (BuffGroupType 190). Each time a skill is used, the accumulator for that skill ID increases by the buff\'s value. The bonus is then applied multiplicatively to skill damage:')
add_code('damage = roundInt(damage * round(1 + round(recordDamage[skillId] / 10000)))')
add_body('For example, if recordDamage = 5000 (50%), then damage is multiplied by (1 + 0.5) = 1.5x.')
add_body('The accumulator persists for the entire battle and is never reset. This means skills used more frequently get progressively stronger over the course of a fight. Applied in BuffSkillValue after Extra Damage and before resistance.')

add_heading('17.2 Counter Damage Multiplier (Skill-Level)', 2)
add_body('counterDamage is a per-skill multiplier that defaults to 1.0 (no effect). It comes from the skill configuration (param5[0]) and is applied to skill damage:')
add_code('damage = roundInt(damage * skill.counterDamage)')
add_body('Despite the name, this is NOT related to counter-attacks — it is a general skill-level multiplier set during skill creation. Applied immediately after recordDamage in the BuffSkillValue pipeline.')

add_heading('17.3 Season PvE Damage Bonus', 2)
add_body('seasonPveDamAdd is a server-controlled bonus that only applies to team 1 (defending team) during season PvE chapters:')
add_code('damage = roundInt(damage * (1 + seasonPveDamAdd))')
add_body('Applied in Unit.addDamage after PvP reduction and before shield absorption. The value comes from the server at battle start and persists for the entire battle. In non-season modes, this value is 0.')

# ============================================================
# SECTION 18: PVP INJURY REDUCE TABLE
# ============================================================
doc.add_page_break()
add_heading('18. PvP Injury Reduce Table (Complete)', 1)

add_body('The complete PvP damage reduction factor for all 220 levels, extracted from Level.json. The raw value is divided by 10,000 to get the actual factor. Damage is divided by this factor, so a factor of 25x means you deal 1/25th of your PvE damage in PvP:')

add_table(
    ['Avg Level', 'Raw Value', 'Factor', 'PvP Damage'],
    [
        ['1', '10,000', '1.0x', '100% of PvE'],
        ['10', '14,000', '1.4x', '~71%'],
        ['20', '34,500', '3.5x', '~29%'],
        ['30', '58,000', '5.8x', '~17%'],
        ['40', '85,000', '8.5x', '~12%'],
        ['50', '110,000', '11.0x', '~9%'],
        ['60', '137,500', '13.8x', '~7%'],
        ['70', '192,000', '19.2x', '~5%'],
        ['80', '276,000', '27.6x', '~4%'],
        ['90', '396,000', '39.6x', '~3%'],
        ['100', '569,000', '56.9x', '~2%'],
        ['120', '1,173,000', '117.3x', '~0.9%'],
        ['140', '2,216,000', '221.6x', '~0.5%'],
        ['160', '3,423,000', '342.3x', '~0.3%'],
        ['180', '4,717,000', '471.7x', '~0.2%'],
        ['200', '6,090,000', '609.0x', '~0.16%'],
        ['220', '7,540,000', '754.0x', '~0.13%'],
    ]
)
add_body('The factor grows roughly exponentially through level 130, then transitions to roughly linear growth (~67-73 per level). At high levels, PvP damage is reduced to a tiny fraction of PvE damage — a level 200 fight deals about 0.16% of PvE damage.')

# ============================================================
# SECTION 19: ATTRIBUTE CAPS
# ============================================================
doc.add_page_break()
add_heading('19. Attribute Caps', 1)

add_body('The following attributes have hard upper limits enforced by the game (from the up_limit field in Attribute.json). When a stat reaches its cap, further bonuses are ignored:')

add_table(
    ['Attribute', 'ID', 'Cap', 'Meaning'],
    [
        ['att_hpsteal', '1014', '100%', 'ATK HP Steal'],
        ['skill_hpsteal', '1015', '100%', 'Skill HP Steal'],
        ['att_resist', '1018', '80%', 'Basic ATK Resistance'],
        ['skill_resist', '1019', '80%', 'Skill Resistance'],
        ['partner_resist', '1020', '80%', 'Pal Resistance'],
        ['resist', '1021', '80%', 'DMG Resistance'],
        ['double_hit_def', '1034', '80%', 'Combo Resistance'],
        ['counter_def', '1035', '80%', 'Counter Resistance'],
        ['control_res', '1042', '100%', 'Control Resistance'],
        ['boss_def', '1052', '80%', 'Boss Defense'],
        ['season_cannon_att_def', '1059', '60%', 'Cannon ATK Defense'],
    ]
)
add_body('All other combat attributes — including ATK, HP, DEF, all damage multipliers, crit rate, crit damage, and Total DMG Bonus/RES — have NO cap. They can be stacked indefinitely.')

# ============================================================
# SECTION 20: EXTENDED ATTRIBUTE SYSTEM
# ============================================================
doc.add_page_break()
add_heading('20. Extended Attribute System (192 Total Attributes)', 1)

add_body('The game has 192 attributes organized across 7 ID ranges. The 82 core battle attributes (IDs 1001-1082) are documented in Section 2. Here are the additional ranges used by the stat assembly pipeline:')

add_table(
    ['ID Range', 'Count', 'Category', 'Examples'],
    [
        ['1-24', '4', 'Base Totals', 'total_att, total_hp, total_att_speed, total_def'],
        ['1001-1082', '82', 'Core Battle', 'att, hp, def, crit_rate, etc. (Section 2)'],
        ['2001-2036', '36', 'Bonus/Add', 'att_base_add, hp_add, def_base_add, crit_dam_add'],
        ['3001-3024', '4', 'Cumulative Totals', 'att_total_add, hp_total_add, def_total_add'],
        ['4001-4006', '6', 'Partner Effects', 'partner_crit_rate, partner_partner_dam, partner_att_speed'],
        ['5001-5012', '12', 'Rogue Mode', 'rogue_att, rogue_def, rogue_hp, rogue_crit_dam'],
        ['6001-6007', '7', 'Spirit System', 'spirit_dam_add, spirit_att, spirit_hp (Section 2)'],
        ['10001-10030', '30', 'Season', 'season_att, season_hp, season_pve_dam_add'],
    ]
)
add_body('The 2000-range attributes are percentage-based group bonuses that feed into the MetaAttrib calculation for their parent core attribute. For example, att_base_add (2001) and att_add (2002) both contribute to att (1001) through the group bonus system documented in Section 2\'s MetaAttrib formula.')
add_body('The 4000-range partner attributes are dedicated modifiers for pal/pet units, while the 5000-range rogue attributes are bonuses specific to the rogue game mode. The 10000-range season attributes are used during seasonal content events (sailing, ships, etc.).')

# ============================================================
# SAVE
# ============================================================
out_path = os.path.expanduser('/home/user/RE13021169/reverse-engineered/LOM_Combat_Reference_Plain_English.docx')
doc.save(out_path)
print(f"Saved to {out_path}")
